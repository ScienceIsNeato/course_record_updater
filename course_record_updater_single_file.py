#!/usr/bin/env python3
"""
Course Record Updater - Single File Version
Consolidates Flask app, database service, adapters, and templates into one executable file.
"""

import datetime
import os
import sys
import tempfile
from dataclasses import dataclass
from typing import Any, Dict, List, Optional

import docx
from flask import (
    Flask,
    flash,
    jsonify,
    redirect,
    render_template_string,
    request,
    url_for,
)
from google.cloud import firestore

# =============================================================================
# CONFIGURATION AND CONSTANTS
# =============================================================================

ALLOWED_EXTENSIONS = {"docx"}
ALLOWED_TERMS = ["FA2024", "SP2024", "SU2024", "FA2025", "SP2025", "SU2025"]
COURSES_COLLECTION = "courses"
UNIQUE_FIELDS = ["term", "course_number", "instructor_name"]


# =============================================================================
# CUSTOM EXCEPTIONS
# =============================================================================


class ValidationError(ValueError):
    pass


class DispatcherError(Exception):
    pass


# =============================================================================
# DATABASE SERVICE
# =============================================================================


class DatabaseService:
    """Handles all Firestore database operations."""

    def __init__(self):
        self.db = None
        self._initialize_firestore()

    def _initialize_firestore(self):
        """Initialize Firestore client."""
        emulator_host = os.environ.get("FIRESTORE_EMULATOR_HOST")

        try:
            self.db = firestore.Client()
            if emulator_host:
                print(
                    f"Firestore client initialized, attempting to connect to emulator at: {emulator_host}"
                )
            else:
                print("Firestore client initialized for cloud connection.")
        except Exception as e:
            print(f"Error initializing Firestore client: {e}")
            if emulator_host:
                print(
                    f"Ensure the Firestore emulator is running and accessible at {emulator_host}"
                )
            self.db = None

    def check_course_exists(self, course_data: Dict) -> Optional[str]:
        """Check if a course with the same unique fields already exists."""
        if not self.db:
            print("[DB Service] check_course_exists: DB unavailable.")
            return None

        if not all(field in course_data for field in UNIQUE_FIELDS):
            print(
                f"[DB Service] check_course_exists: Missing unique fields ({UNIQUE_FIELDS})."
            )
            return None

        try:
            query = self.db.collection(COURSES_COLLECTION)
            for field in UNIQUE_FIELDS:
                query = query.where(
                    filter=firestore.FieldFilter(field, "==", course_data[field])
                )

            docs = query.limit(1).stream()
            first_doc = next(docs, None)

            if first_doc:
                print(f"[DB Service] Found existing course with ID: {first_doc.id}")
                return first_doc.id
            else:
                print("[DB Service] No existing course found.")
                return None

        except Exception as e:
            print(f"[DB Service] Error querying Firestore: {e}")
            return None

    def save_course(self, course_data: Dict) -> Optional[str]:
        """Save course data to Firestore, checking for duplicates first."""
        print("[DB Service] save_course called.")
        if not self.db:
            print("[DB Service] Firestore client not available.")
            return None

        if not isinstance(course_data, dict):
            print("[DB Service] Invalid data type provided.")
            return None

        # Check for duplicates
        existing_id = self.check_course_exists(course_data)
        if existing_id:
            print(f"[DB Service] Duplicate detected. Existing ID: {existing_id}")
            return f"DUPLICATE:{existing_id}"

        # Save new course
        try:
            data_to_save = course_data.copy()
            data_to_save["timestamp"] = firestore.SERVER_TIMESTAMP
            collection_ref = self.db.collection(COURSES_COLLECTION)
            _, doc_ref = collection_ref.add(data_to_save)
            print(f"[DB Service] Course saved with ID: {doc_ref.id}")
            return doc_ref.id
        except Exception as e:
            print(f"[DB Service] Error saving course: {e}")
            return None

    def get_all_courses(self) -> List[Dict]:
        """Retrieve all courses from Firestore."""
        print("[DB Service] get_all_courses called.")
        if not self.db:
            print("[DB Service] Firestore client not available.")
            return []

        try:
            courses_ref = self.db.collection(COURSES_COLLECTION)
            query = courses_ref.order_by(
                "timestamp", direction=firestore.Query.DESCENDING
            )
            docs = query.stream()

            courses_list = []
            for doc in docs:
                course = doc.to_dict()
                course["id"] = doc.id
                courses_list.append(course)

            print(f"[DB Service] Retrieved {len(courses_list)} courses.")
            return courses_list
        except Exception as e:
            print(f"[DB Service] Error getting courses: {e}")
            return []

    def update_course(self, course_id: str, update_data: Dict) -> bool:
        """Update existing course document."""
        print(f"[DB Service] update_course called for ID: {course_id}")
        if not self.db:
            print("[DB Service] Firestore client not available.")
            return False

        if not course_id or not isinstance(update_data, dict):
            print("[DB Service] Invalid arguments.")
            return False

        try:
            doc_ref = self.db.collection(COURSES_COLLECTION).document(course_id)
            update_data_with_ts = update_data.copy()
            update_data_with_ts["last_modified"] = firestore.SERVER_TIMESTAMP
            doc_ref.update(update_data_with_ts)
            print(f"[DB Service] Course updated for ID: {course_id}")
            return True
        except Exception as e:
            print(f"[DB Service] Error updating course {course_id}: {e}")
            return False

    def delete_course(self, course_id: str) -> bool:
        """Delete course document from Firestore."""
        print(f"[DB Service] delete_course called for ID: {course_id}")
        if not self.db:
            print("[DB Service] Firestore client not available.")
            return False

        if not course_id:
            print("[DB Service] Invalid course_id.")
            return False

        try:
            doc_ref = self.db.collection(COURSES_COLLECTION).document(course_id)
            doc_ref.delete()
            print(f"[DB Service] Course deleted for ID: {course_id}")
            return True
        except Exception as e:
            print(f"[DB Service] Error deleting course {course_id}: {e}")
            return False


# =============================================================================
# VALIDATION AND ADAPTERS
# =============================================================================


class BaseAdapter:
    """Base adapter for parsing and validating input data."""

    EXPECTED_FIELDS = {
        "course_title": (True, str, None),
        "course_number": (True, str, None),
        "term": (True, str, lambda t: t in ALLOWED_TERMS),
        "instructor_name": (True, str, None),
        "num_students": (False, int, lambda n: n >= 0),
        "grade_a": (False, int, lambda g: g >= 0),
        "grade_b": (False, int, lambda g: g >= 0),
        "grade_c": (False, int, lambda g: g >= 0),
        "grade_d": (False, int, lambda g: g >= 0),
        "grade_f": (False, int, lambda g: g >= 0),
    }

    GRADE_FIELDS = ["grade_a", "grade_b", "grade_c", "grade_d", "grade_f"]

    def parse_and_validate(self, form_data: Dict) -> Dict:
        """Parse and validate data from form dictionary."""
        validated_data = {}
        errors = []
        raw_input_data = {k: str(v).strip() for k, v in form_data.items()}

        # Check for missing required fields
        for field, config in self.EXPECTED_FIELDS.items():
            is_required = config[0]
            if is_required and not raw_input_data.get(field):
                errors.append(f"Missing required field: {field}")

        # Grade distribution pre-check
        any_grade_entered = any(
            raw_input_data.get(g_field) for g_field in self.GRADE_FIELDS
        )
        num_students_raw = raw_input_data.get("num_students")

        if any_grade_entered and not num_students_raw:
            errors.append(
                "Number of students is required when entering grade distribution."
            )

        if errors:
            raise ValidationError("; ".join(errors))

        # Process fields
        grade_values = {}
        parsed_num_students = None

        for field, value_str in raw_input_data.items():
            if field in self.EXPECTED_FIELDS:
                config = self.EXPECTED_FIELDS[field]
                is_required_field = config[0]
                expected_type = config[1]
                validator = config[2]

                is_num_students_and_grades_entered = (
                    field == "num_students" and any_grade_entered
                )
                if (
                    not value_str
                    and not is_required_field
                    and not is_num_students_and_grades_entered
                ):
                    continue

                # Type conversion
                processed_value = None
                conversion_error = False
                if expected_type:
                    try:
                        if expected_type == int and value_str:
                            processed_value = int(value_str)
                        elif expected_type == float and value_str:
                            processed_value = float(value_str)
                        elif expected_type == str:
                            processed_value = value_str
                        else:
                            processed_value = value_str
                    except (ValueError, TypeError):
                        errors.append(
                            f"Invalid value for {field}: Cannot convert '{value_str}' to {expected_type.__name__}"
                        )
                        conversion_error = True
                else:
                    processed_value = value_str

                if conversion_error:
                    continue

                # Store values for cross-field validation
                if field == "num_students" and processed_value is not None:
                    parsed_num_students = processed_value
                elif field in self.GRADE_FIELDS and processed_value is not None:
                    grade_values[field] = processed_value

                # Run validator
                if validator and processed_value is not None:
                    try:
                        if not validator(processed_value):
                            errors.append(
                                f"Invalid value for {field}: Failed validation rule."
                            )
                    except Exception as e:
                        errors.append(f"Error during validation for {field}: {e}")

                # Store validated field
                if field not in [e.split(":")[0].split(" ")[-1] for e in errors]:
                    if processed_value is not None or is_required_field:
                        validated_data[field] = processed_value

        # Cross-field validation (grade sum)
        if any_grade_entered:
            if parsed_num_students is not None:
                current_grade_sum = sum(grade_values.values())
                if current_grade_sum != parsed_num_students:
                    errors.append(
                        f"Sum of grades ({current_grade_sum}) does not match Number of Students ({parsed_num_students})."
                    )

        if errors:
            raise ValidationError("; ".join(errors))

        # Remove empty grade fields
        if not any_grade_entered:
            for g_field in self.GRADE_FIELDS:
                validated_data.pop(g_field, None)

        return validated_data


class FileAdapterDispatcher:
    """Dispatcher for file-based adapters."""

    def __init__(self, use_base_validation=True):
        self.use_base_validation = use_base_validation
        self.base_adapter = BaseAdapter() if use_base_validation else None
        self.adapters = {
            "dummy_adapter": self._dummy_parse,
            "nursing_sample_adapter": self._nursing_parse,
            "business_sample_adapter": self._business_parse,
        }

    def process_file(self, document, adapter_name: str) -> List[Dict]:
        """Process file using specified adapter."""
        if adapter_name not in self.adapters:
            raise DispatcherError(f"Unknown adapter: {adapter_name}")

        parse_func = self.adapters[adapter_name]
        raw_data_list = parse_func(document)

        if not isinstance(raw_data_list, list):
            raw_data_list = [raw_data_list]

        validated_list = []
        for raw_data in raw_data_list:
            if self.use_base_validation:
                try:
                    validated_data = self.base_adapter.parse_and_validate(raw_data)
                    validated_list.append(validated_data)
                except ValidationError as e:
                    print(f"Validation error for record: {e}")
                    continue
            else:
                validated_list.append(raw_data)

        return validated_list

    def _dummy_parse(self, document) -> Dict:
        """Dummy parser for testing."""
        return {
            "course_title": "Sample Course from Dummy",
            "course_number": "DUMMY101",
            "term": "FA2024",
            "instructor_name": "Dr. Test",
            "num_students": "25",
            "grade_a": "5",
            "grade_b": "10",
            "grade_c": "8",
            "grade_d": "2",
            "grade_f": "0",
        }

    def _nursing_parse(self, document) -> List[Dict]:
        """Parser for nursing sample format."""
        courses = []
        current_course = {}

        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue

            if text.startswith("Course:"):
                if current_course:
                    courses.append(current_course)
                current_course = {"course_title": text.replace("Course:", "").strip()}
            elif text.startswith("Number:"):
                current_course["course_number"] = text.replace("Number:", "").strip()
            elif text.startswith("Term:"):
                current_course["term"] = text.replace("Term:", "").strip()
            elif text.startswith("Instructor:"):
                current_course["instructor_name"] = text.replace(
                    "Instructor:", ""
                ).strip()
            elif text.startswith("Students:"):
                current_course["num_students"] = text.replace("Students:", "").strip()
            elif text.startswith("Grades:"):
                grades_text = text.replace("Grades:", "").strip()
                try:
                    grades = [int(x.strip()) for x in grades_text.split(",")]
                    if len(grades) == 5:
                        current_course.update(
                            {
                                "grade_a": str(grades[0]),
                                "grade_b": str(grades[1]),
                                "grade_c": str(grades[2]),
                                "grade_d": str(grades[3]),
                                "grade_f": str(grades[4]),
                            }
                        )
                except (ValueError, IndexError):
                    print(f"Could not parse grades: {grades_text}")

        if current_course:
            courses.append(current_course)

        return courses

    def _business_parse(self, document) -> List[Dict]:
        """Parser for business sample format."""
        courses = []

        for table in document.tables:
            if len(table.rows) < 2:
                continue

            headers = [cell.text.strip().lower() for cell in table.rows[0].cells]

            for row in table.rows[1:]:
                if len(row.cells) < len(headers):
                    continue

                course_data = {}
                for i, cell in enumerate(row.cells):
                    if i < len(headers):
                        header = headers[i]
                        value = cell.text.strip()

                        if "course" in header and "title" in header:
                            course_data["course_title"] = value
                        elif "course" in header and "number" in header:
                            course_data["course_number"] = value
                        elif "term" in header:
                            course_data["term"] = value
                        elif "instructor" in header:
                            course_data["instructor_name"] = value
                        elif "students" in header:
                            course_data["num_students"] = value
                        elif "grade a" in header:
                            course_data["grade_a"] = value
                        elif "grade b" in header:
                            course_data["grade_b"] = value
                        elif "grade c" in header:
                            course_data["grade_c"] = value
                        elif "grade d" in header:
                            course_data["grade_d"] = value
                        elif "grade f" in header:
                            course_data["grade_f"] = value

                if course_data:
                    courses.append(course_data)

        return courses


# =============================================================================
# HTML TEMPLATE
# =============================================================================

HTML_TEMPLATE = """
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>CEI Course Admin</title>
    <link href="https://cdn.jsdelivr.net/npm/bootstrap@5.3.0/dist/css/bootstrap.min.css" rel="stylesheet">
    <style>
        .flash-message { margin-top: 1rem; }
        .navbar-brand img { max-height: 40px; margin-right: 10px; }
        .editable { cursor: pointer; }
        .editable:hover { background-color: #f8f9fa; }
    </style>
</head>
<body>
    <nav class="navbar navbar-expand-lg navbar-dark bg-dark mb-4">
        <div class="container-fluid">
            <a class="navbar-brand d-flex align-items-center" href="/">
                CEI Course Admin Tool
            </a>
        </div>
    </nav>

    <div class="container mt-4">
        <!-- Display Upload Results -->
        {% if upload_results %}
        <div class="card mb-4 border-secondary">
            <div class="card-header bg-secondary text-white">Upload Results</div>
            <div class="card-body">
                {% if upload_results.success %}
                    <div class="alert alert-success mb-2" role="alert">
                        <strong>Successfully Saved:</strong>
                        <ul>
                            {% for msg in upload_results.success %}<li>{{ msg }}</li>{% endfor %}
                        </ul>
                    </div>
                {% endif %}
                {% if upload_results.duplicate %}
                    <div class="alert alert-warning mb-2" role="alert">
                        <strong>Skipped Duplicates:</strong>
                        <ul>
                            {% for msg in upload_results.duplicate %}<li>{{ msg }}</li>{% endfor %}
                        </ul>
                    </div>
                {% endif %}
                {% if upload_results.failed %}
                    <div class="alert alert-danger mb-0" role="alert">
                        <strong>Failed to Save:</strong>
                        <ul>
                            {% for msg in upload_results.failed %}<li>{{ msg }}</li>{% endfor %}
                        </ul>
                    </div>
                {% endif %}
            </div>
        </div>
        {% endif %}

        <!-- Flash Messages -->
        {% with messages = get_flashed_messages(with_categories=true) %}
            {% if messages %}
                {% for category, message in messages %}
                    <div class="alert alert-{{ category if category != 'message' else 'info' }} alert-dismissible fade show flash-message" role="alert">
                        {{ message | safe }}
                        <button type="button" class="btn-close" data-bs-dismiss="alert" aria-label="Close"></button>
                    </div>
                {% endfor %}
            {% endif %}
        {% endwith %}

        <!-- File Upload Form -->
        <div class="card mb-4">
            <div class="card-header">Upload Course Document (.docx)</div>
            <div class="card-body">
                <form action="{{ url_for('add_course_automatic') }}" method="post" enctype="multipart/form-data">
                    <div class="row g-3 align-items-end">
                        <div class="col-md-6">
                            <label for="course_document" class="form-label">Select Word Document:</label>
                            <input type="file" class="form-control" id="course_document" name="course_document" accept=".docx" required>
                        </div>
                        <div class="col-md-4">
                            <label for="adapter_name" class="form-label">Select Document Format:</label>
                            <select class="form-select" id="adapter_name" name="adapter_name" required>
                                <option value="" disabled selected>-- Choose Format --</option>
                                <option value="nursing_sample_adapter">Nursing Format Sample</option>
                                <option value="business_sample_adapter">Business Format Sample</option>
                                <option value="dummy_adapter">Dummy Format (Testing)</option>
                            </select>
                        </div>
                        <div class="col-md-2">
                            <button type="submit" class="btn btn-secondary w-100">Upload</button>
                        </div>
                    </div>
                </form>
            </div>
        </div>

        <!-- Manual Entry Form -->
        <div class="card mb-4">
            <div class="card-header">Add New Course Manually</div>
            <div class="card-body">
                <form action="{{ url_for('add_course') }}" method="post">
                    <div class="row g-3">
                        <div class="col-md-6">
                            <label for="course_number" class="form-label">Course Number*</label>
                            <input type="text" class="form-control" id="course_number" name="course_number" value="{{ form_data.get('course_number', '') }}" required>
                        </div>
                        <div class="col-md-6">
                            <label for="course_title" class="form-label">Course Title*</label>
                            <input type="text" class="form-control" id="course_title" name="course_title" value="{{ form_data.get('course_title', '') }}" required>
                        </div>
                        <div class="col-md-4">
                            <label for="num_students" class="form-label">Number of Students</label>
                            <input type="number" class="form-control" id="num_students" name="num_students" value="{{ form_data.get('num_students', '') }}" min="0">
                        </div>
                        <div class="col-md-4">
                            <label for="instructor_name" class="form-label">Instructor Name*</label>
                            <input type="text" class="form-control" id="instructor_name" name="instructor_name" value="{{ form_data.get('instructor_name', '') }}" required>
                        </div>
                        <div class="col-md-4">
                            <label for="term" class="form-label">Term*</label>
                            <select class="form-select" id="term" name="term" required>
                                <option value="" disabled selected>Select Term</option>
                                {% for term_option in allowed_terms %}
                                <option value="{{ term_option }}" {% if form_data.get('term') == term_option %}selected{% endif %}>{{ term_option }}</option>
                                {% endfor %}
                            </select>
                        </div>
                        
                        <h5 class="mt-4">Grade Distribution (Optional)</h5>
                        <div class="col-md-2">
                            <label for="grade_a" class="form-label">Grade A</label>
                            <input type="number" class="form-control" id="grade_a" name="grade_a" value="{{ form_data.get('grade_a', '') }}" min="0">
                        </div>
                        <div class="col-md-2">
                            <label for="grade_b" class="form-label">Grade B</label>
                            <input type="number" class="form-control" id="grade_b" name="grade_b" value="{{ form_data.get('grade_b', '') }}" min="0">
                        </div>
                        <div class="col-md-2">
                            <label for="grade_c" class="form-label">Grade C</label>
                            <input type="number" class="form-control" id="grade_c" name="grade_c" value="{{ form_data.get('grade_c', '') }}" min="0">
                        </div>
                        <div class="col-md-2">
                            <label for="grade_d" class="form-label">Grade D</label>
                            <input type="number" class="form-control" id="grade_d" name="grade_d" value="{{ form_data.get('grade_d', '') }}" min="0">
                        </div>
                        <div class="col-md-2">
                            <label for="grade_f" class="form-label">Grade F</label>
                            <input type="number" class="form-control" id="grade_f" name="grade_f" value="{{ form_data.get('grade_f', '') }}" min="0">
                        </div>
                    </div>
                    <button type="submit" class="btn btn-primary mt-3">Add Course</button>
                </form>
            </div>
        </div>

        <!-- Course List -->
        <div class="card">
            <div class="card-header">Existing Courses</div>
            <div class="card-body">
                <table class="table table-striped">
                    <thead>
                        <tr>
                            <th>Course Number</th>
                            <th>Title</th>
                            <th>Instructor</th>
                            <th>Term</th>
                            <th># Students</th>
                            <th>Grade A</th>
                            <th>Grade B</th>
                            <th>Grade C</th>
                            <th>Grade D</th>
                            <th>Grade F</th>
                            <th>Actions</th>
                        </tr>
                    </thead>
                    <tbody>
                        {% for course in courses %}
                        <tr data-course-id="{{ course.id }}">
                            <td class="editable" data-field="course_number">{{ course.course_number }}</td>
                            <td class="editable" data-field="course_title">{{ course.course_title }}</td>
                            <td class="editable" data-field="instructor_name">{{ course.instructor_name }}</td>
                            <td class="editable" data-field="term">{{ course.term }}</td>
                            <td class="editable" data-field="num_students">{{ course.num_students | default('N/A', true) }}</td>
                            <td class="editable" data-field="grade_a">{{ course.grade_a | default('N/A', true) }}</td>
                            <td class="editable" data-field="grade_b">{{ course.grade_b | default('N/A', true) }}</td>
                            <td class="editable" data-field="grade_c">{{ course.grade_c | default('N/A', true) }}</td>
                            <td class="editable" data-field="grade_d">{{ course.grade_d | default('N/A', true) }}</td>
                            <td class="editable" data-field="grade_f">{{ course.grade_f | default('N/A', true) }}</td>
                            <td>
                                <button class="btn btn-sm btn-outline-primary edit-btn" data-course-id="{{ course.id }}">Edit</button>
                                <button class="btn btn-sm btn-outline-danger delete-btn" data-course-id="{{ course.id }}">Delete</button>
                            </td>
                        </tr>
                        {% endfor %}
                    </tbody>
                </table>
            </div>
        </div>
    </div>

    <script src="https://cdn.jsdelivr.net/npm/bootstrap@5.3.0/dist/js/bootstrap.bundle.min.js"></script>
    <script>
        // Basic inline editing functionality
        document.addEventListener('DOMContentLoaded', function() {
            // Edit functionality
            document.querySelectorAll('.edit-btn').forEach(btn => {
                btn.addEventListener('click', function() {
                    const courseId = this.getAttribute('data-course-id');
                    const row = document.querySelector(`tr[data-course-id="${courseId}"]`);
                    
                    if (this.textContent === 'Edit') {
                        // Enable editing
                        row.querySelectorAll('.editable').forEach(cell => {
                            const field = cell.getAttribute('data-field');
                            const currentValue = cell.textContent === 'N/A' ? '' : cell.textContent;
                            cell.innerHTML = `<input type="text" class="form-control form-control-sm" value="${currentValue}" data-field="${field}">`;
                        });
                        this.textContent = 'Save';
                        this.className = 'btn btn-sm btn-success edit-btn';
                    } else {
                        // Save changes
                        const updateData = {};
                        row.querySelectorAll('.editable input').forEach(input => {
                            const field = input.getAttribute('data-field');
                            updateData[field] = input.value || null;
                        });
                        
                        fetch(`/edit_course/${courseId}`, {
                            method: 'PUT',
                            headers: {'Content-Type': 'application/json'},
                            body: JSON.stringify(updateData)
                        })
                        .then(response => response.json())
                        .then(data => {
                            if (data.success) {
                                location.reload();
                            } else {
                                alert('Error updating course: ' + data.error);
                            }
                        })
                        .catch(error => {
                            alert('Error updating course: ' + error);
                        });
                    }
                });
            });
            
            // Delete functionality
            document.querySelectorAll('.delete-btn').forEach(btn => {
                btn.addEventListener('click', function() {
                    const courseId = this.getAttribute('data-course-id');
                    
                    if (confirm('Are you sure you want to delete this course?')) {
                        fetch(`/delete_course/${courseId}`, {
                            method: 'DELETE',
                            headers: {'Content-Type': 'application/json'}
                        })
                        .then(response => response.json())
                        .then(data => {
                            if (data.success) {
                                location.reload();
                            } else {
                                alert('Error deleting course: ' + data.error);
                            }
                        })
                        .catch(error => {
                            alert('Error deleting course: ' + error);
                        });
                    }
                });
            });
        });
    </script>
</body>
</html>
"""


# =============================================================================
# FLASK APPLICATION
# =============================================================================

app = Flask(__name__)
app.secret_key = os.environ.get("FLASK_SECRET_KEY", "dev-secret-key")

# Initialize services
db_service = DatabaseService()
form_adapter = BaseAdapter()


def allowed_file(filename):
    """Check if filename has allowed extension."""
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS


@app.route("/")
def index():
    """Render main page with course list and entry form."""
    courses = []
    if db_service.db is None:
        flash(
            "Error: Database service is not available. Check Firestore connection.",
            "error",
        )
    else:
        try:
            courses = db_service.get_all_courses()
        except Exception as e:
            flash(f"Error fetching courses: {e}", "error")

    return render_template_string(
        HTML_TEMPLATE, courses=courses, allowed_terms=ALLOWED_TERMS, form_data={}
    )


@app.route("/add", methods=["POST"])
def add_course():
    """Handle manual course form submission."""
    if db_service.db is None:
        flash("Error: Cannot add course, database service is unavailable.", "error")
        return redirect(url_for("index"))

    form_data = request.form.to_dict()
    try:
        validated_data = form_adapter.parse_and_validate(form_data)

        # Clean up optional None fields
        if "num_students" in validated_data and validated_data["num_students"] is None:
            del validated_data["num_students"]
        for grade_field in BaseAdapter.GRADE_FIELDS:
            if grade_field in validated_data and validated_data[grade_field] is None:
                del validated_data[grade_field]

        course_id = db_service.save_course(validated_data)
        if course_id and not course_id.startswith("DUPLICATE:"):
            flash(
                f"Course '{validated_data['course_title']}' added successfully with ID: {course_id}",
                "success",
            )
        elif course_id.startswith("DUPLICATE:"):
            flash("Course already exists (duplicate detected).", "warning")
        else:
            flash("Error: Failed to save course to the database.", "error")

    except ValidationError as e:
        error_list = str(e).split("; ")
        formatted_message = "Please fix the following issues:<br><ul>"
        for error in error_list:
            formatted_message += f"<li>{error}</li>"
        formatted_message += "</ul>"
        flash(formatted_message, "danger")

        try:
            courses = db_service.get_all_courses()
        except Exception:
            courses = []

        return render_template_string(
            HTML_TEMPLATE,
            courses=courses,
            allowed_terms=ALLOWED_TERMS,
            form_data=form_data,
        )
    except Exception as e:
        flash(f"An error occurred while adding the course: {e}", "error")

    return redirect(url_for("index"))


@app.route("/add_course_automatic", methods=["POST"])
def add_course_automatic():
    """Handle file upload and automatic processing."""
    selected_adapter_name = request.form.get("adapter_name")
    courses = []
    upload_results = {"success": [], "duplicate": [], "failed": []}

    if db_service.db is None:
        flash("Error: Database service unavailable, cannot process upload.", "error")
        return redirect(url_for("index"))

    try:
        courses = db_service.get_all_courses()
    except Exception as fetch_e:
        flash(f"Error fetching courses: {fetch_e}", "error")

    file = request.files.get("course_document")

    # Basic validation
    error_message = None
    if not file or file.filename == "":
        error_message = "No file selected."
    elif not selected_adapter_name:
        error_message = "No document format/adapter selected."
    elif not allowed_file(file.filename):
        error_message = (
            f"File type not allowed ('{file.filename}'). Only .docx is permitted."
        )

    if error_message:
        flash(
            f"Please fix the following issues:<br><ul><li>{error_message}</li></ul>",
            "danger",
        )
        return render_template_string(
            HTML_TEMPLATE,
            courses=courses,
            allowed_terms=ALLOWED_TERMS,
            form_data={},
            selected_adapter_name=selected_adapter_name,
            upload_results=None,
        )

    # Process file
    try:
        print(
            f"Processing uploaded file: {file.filename} with adapter: {selected_adapter_name}"
        )
        document = docx.Document(file.stream)

        dispatcher = FileAdapterDispatcher(use_base_validation=True)
        validated_data_list = dispatcher.process_file(document, selected_adapter_name)

        if not validated_data_list:
            flash(
                f"Warning: Adapter '{selected_adapter_name}' parsed 0 records from {file.filename}.",
                "warning",
            )
            return render_template_string(
                HTML_TEMPLATE,
                courses=courses,
                allowed_terms=ALLOWED_TERMS,
                form_data={},
                selected_adapter_name=selected_adapter_name,
                upload_results=upload_results,
            )

        # Save multiple courses
        for course_data in validated_data_list:
            course_identifier = f"{course_data.get('course_number', 'N/A')} ({course_data.get('term', 'N/A')})"
            save_result = db_service.save_course(course_data)

            if isinstance(save_result, str) and save_result.startswith("DUPLICATE:"):
                upload_results["duplicate"].append(
                    f"{course_identifier} - Already exists"
                )
            elif save_result:
                upload_results["success"].append(
                    f"{course_identifier} - Saved with ID: {save_result}"
                )
            else:
                upload_results["failed"].append(f"{course_identifier} - Database error")

        # Update courses list
        courses = db_service.get_all_courses()

    except Exception as e:
        flash(f"Error processing file: {e}", "error")
        return redirect(url_for("index"))

    return render_template_string(
        HTML_TEMPLATE,
        courses=courses,
        allowed_terms=ALLOWED_TERMS,
        form_data={},
        selected_adapter_name=selected_adapter_name,
        upload_results=upload_results,
    )


@app.route("/edit_course/<string:course_id>", methods=["PUT", "POST"])
def edit_course(course_id):
    """Handle updating existing course."""
    if db_service.db is None:
        return jsonify({"error": "Database service unavailable"}), 503

    if not request.is_json:
        return jsonify({"error": "Request must be JSON"}), 400

    update_data = request.get_json()
    if not update_data:
        return jsonify({"error": "No data provided"}), 400

    if db_service.update_course(course_id, update_data):
        return jsonify({"success": True, "message": "Course updated successfully"}), 200
    else:
        return jsonify({"error": "Failed to update course"}), 500


@app.route("/delete_course/<string:course_id>", methods=["DELETE", "POST"])
def delete_course(course_id):
    """Handle deleting course."""
    if db_service.db is None:
        return jsonify({"error": "Database service unavailable"}), 503

    if db_service.delete_course(course_id):
        return jsonify({"success": True, "message": "Course deleted successfully"}), 200
    else:
        return jsonify({"error": "Failed to delete course"}), 500


# =============================================================================
# MAIN EXECUTION
# =============================================================================

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 8080))
    debug = os.environ.get("FLASK_DEBUG", "False").lower() == "true"

    print("=" * 60)
    print("CEI Course Record Updater - Single File Version")
    print("=" * 60)
    print(f"Starting Flask application on port {port}")
    print(f"Debug mode: {debug}")
    print(
        f"Firestore emulator: {'Yes' if os.environ.get('FIRESTORE_EMULATOR_HOST') else 'No'}"
    )
    print("=" * 60)

    app.run(host="0.0.0.0", port=port, debug=debug)
