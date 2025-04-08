# generate_sample_docx.py
import sys
try:
    from docx import Document
    from docx.shared import Pt
except ImportError:
    print("Error: python-docx library not found.")
    print("Please install it using: pip install python-docx")
    sys.exit(1)

# --- Data ---

nursing_data = [
    {"course_number": "NUR-101", "course_title": "Fundamentals of Nursing", "instructor_name": "Prof. Nightingale", "term": "FA2024", "num_students": "45", "grade_a": "10", "grade_b": "20", "grade_c": "10", "grade_d": "3", "grade_f": "2"},
    {"course_number": "NUR-105", "course_title": "Medical Terminology", "instructor_name": "Prof. Caduceus", "term": "FA2024", "num_students": "50", "grade_a": "", "grade_b": "", "grade_c": "", "grade_d": "", "grade_f": ""},
    {"course_number": "NUR-210", "course_title": "Pharmacology for Nurses", "instructor_name": "Prof. Nightingale", "term": "SP2025", "num_students": "40", "grade_a": "15", "grade_b": "15", "grade_c": "8", "grade_d": "2", "grade_f": "0"},
]

business_data = [
    {"course_number": "BUS-101", "course_title": "Introduction to Business", "instructor_name": "Prof. Smith", "term": "FA2024", "num_students": "30", "grades_str": "A=5, B=15, C=8, D=1, F=1"},
    {"course_number": "ACCT-110", "course_title": "Principles of Accounting I", "instructor_name": "Prof. Ledger", "term": "FA2024", "num_students": "25", "grades_str": "A=4, B=10, C=6, D=3, F=2"},
    {"course_number": "MKTG-201", "course_title": "Marketing Principles", "instructor_name": "Prof. Adwell", "term": "SP2025", "num_students": "35", "grades_str": "N/A"},
    {"course_number": "BUS-250", "course_title": "Business Law", "instructor_name": "Prof. Gavel", "term": "SP2025", "num_students": "28", "grades_str": "A=7, B=12, C=5, D=4, F=0"},
    {"course_number": "ECON-201", "course_title": "Principles of Macroeconomics", "instructor_name": "Prof. Keynes", "term": "FA2024", "num_students": "40", "grades_str": "A=8, B=18, C=10, D=3, F=1"},
]

# --- Nursing Docx (Table Format) ---

def create_nursing_docx(filename="nursing_sample.docx"):
    print(f"Creating {filename}...")
    document = Document()
    document.add_heading('Nursing Program Courses - Sample', level=1)

    headers = ["Course Number", "Title", "Instructor", "Term", "Students", "Grade A", "Grade B", "Grade C", "Grade D", "Grade F"]
    table = document.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid' # Apply basic grid style

    # Populate header row
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        # Optional: Make header bold
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True

    # Populate data rows
    for course in nursing_data:
        row_cells = table.add_row().cells
        row_cells[0].text = course.get("course_number", "")
        row_cells[1].text = course.get("course_title", "")
        row_cells[2].text = course.get("instructor_name", "")
        row_cells[3].text = course.get("term", "")
        row_cells[4].text = course.get("num_students", "")
        row_cells[5].text = course.get("grade_a", "")
        row_cells[6].text = course.get("grade_b", "")
        row_cells[7].text = course.get("grade_c", "")
        row_cells[8].text = course.get("grade_d", "")
        row_cells[9].text = course.get("grade_f", "")

    document.save(filename)
    print(f"{filename} created successfully.")

# --- Business Docx (Text Block Format) ---

def create_business_docx(filename="business_sample.docx"):
    print(f"Creating {filename}...")
    document = Document()
    document.add_heading('Business Program Courses - Sample', level=1)

    for i, course in enumerate(business_data):
        document.add_paragraph("COURSE DATA", style='Heading 2') # Use a style for visual separation
        document.add_paragraph(f"Number: {course.get('course_number', 'N/A')}")
        document.add_paragraph(f"Title: {course.get('course_title', 'N/A')}")
        document.add_paragraph(f"Instructor: {course.get('instructor_name', 'N/A')}")
        document.add_paragraph(f"Term: {course.get('term', 'N/A')}")
        document.add_paragraph(f"Student Count: {course.get('num_students', 'N/A')}")
        document.add_paragraph(f"Grades: {course.get('grades_str', 'N/A')}")

        # Add a separator, except after the last course
        if i < len(business_data) - 1:
            document.add_paragraph("---") # Simple text separator
            # Or use document.add_page_break() or add more spacing if needed

    document.save(filename)
    print(f"{filename} created successfully.")

# --- Main Execution ---

if __name__ == "__main__":
    create_nursing_docx()
    print("-" * 20)
    create_business_docx() 